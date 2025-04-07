import os
import zipfile
import pdfplumber
import re
import time
from docx import Document
import pytesseract
from pdf2image import convert_from_path
from PIL import Image
import subprocess
import yaml
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime
from dateutil.relativedelta import relativedelta

def timed_function(func):
    """Decorator to measure function execution time."""
    def wrapper(*args, **kwargs):
        start_time = time.time()
        result = func(*args, **kwargs)
        end_time = time.time()
        elapsed_time = end_time - start_time
        if elapsed_time > 2:
            print(f"⏱ {func.__name__} took {elapsed_time:.4f} seconds")
        return result
    return wrapper

@timed_function
def load_yaml(yaml_path):
    """Load YAML configuration and return structured data."""
    with open(yaml_path, "r", encoding="utf-8") as file:
        yaml_data = yaml.safe_load(file)
    return yaml_data

@timed_function
def clean_text(text):
    """Clean text while preserving line breaks."""
    lines = text.splitlines()
    cleaned_lines = []
    for line in lines:
        cleaned_line = re.sub(r'[^a-zA-Z0-9\s\n*()\-,.:;?!\'"]', '', line)
        cleaned_lines.append(cleaned_line)
    return "\n".join(cleaned_lines)

@timed_function
def extract_text_from_pdf(pdf_path):
    """Extract text from a PDF, using pdftotext first, then pdfplumber, then OCR if needed."""
    output_dir = "work_files"
    os.makedirs(output_dir, exist_ok=True)
    output_file_path = os.path.join(output_dir, os.path.basename(pdf_path) + ".txt")

    # Try using pdftotext first
    result = subprocess.run(['pdftotext', pdf_path, '-'], capture_output=True, text=True)
    text = result.stdout.strip()

    if text:
        print(f"✅ Extracted text using pdftotext: {text[:100]}...")
        with open(output_file_path, "w", encoding="utf-8") as f:
            f.write(text)
        return text

    print("⚠️ pdftotext failed, trying pdfplumber...")

    # Fallback to pdfplumber
    text = ""
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"

    if text:
        print(f"✅ Extracted text using pdfplumber: {text[:100]}...")
        text = text.strip()
        with open(output_file_path, "w", encoding="utf-8") as f:
            f.write(text)
        return text

    print("⚠️ pdfplumber failed, performing OCR...")

    # Final fallback: Use OCR (slow)
    text = ""
    images = convert_from_path(pdf_path)
    for img in images:
        ocr_text = pytesseract.image_to_string(img, lang='eng', config='--oem 3 --psm 6')
        cleaned_text = ocr_text.strip()
        text += cleaned_text + "\n"

    text = text.strip()
    print(f"✅ Extracted text using OCR (cleaned): {text[:100]}...")

    with open(output_file_path, "w", encoding="utf-8") as f:
        f.write(text)

    return text

@timed_function
def extract_text_from_docx(docx_path):
    """Extract text from Word document with error handling."""
    try:
        doc = Document(docx_path)
        return "\n".join(para.text for para in doc.paragraphs) or ""
    except Exception as e:
        print(f"⚠️ Error extracting text from {docx_path}: {str(e)}")
        return ""

@timed_function
def add_formatted_paragraph(doc, text, style=None, bold=False, italic=False):
    """Add formatted paragraph to document."""
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p

@timed_function
def write_combined_text(text, filename="combined_text.txt"):
    """Write combined text to a file for later processing."""
    output_dir = "work_files"
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, filename)
    
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(text)
    print(f"✅ Combined text saved to {output_path}")
    return output_path

@timed_function
def read_combined_text(filename="combined_text.txt"):
    """Read combined text from file."""
    input_path = os.path.join("work_files", filename)
    try:
        with open(input_path, "r", encoding="utf-8") as f:
            return f.read()
    except FileNotFoundError:
        print(f"⚠️ Combined text file not found: {input_path}")
        return None

@timed_function
def extract_matching_text(text, search_pattern, extract_pattern, message_template):
    """Extract and format text using combined search and extract patterns."""
    try:
        # First find the section using search_pattern
        section_match = re.search(search_pattern, text, re.IGNORECASE | re.DOTALL)
        if not section_match:
            return None
            
        # Then extract the specific content using extract_pattern
        content_match = re.search(extract_pattern, text[section_match.start():], re.IGNORECASE | re.DOTALL)
        if not content_match:
            return None
            
        # Format the extracted content
        extracted = {f"extracted_text_{i+1}": content_match.group(i+1) 
                    for i in range(content_match.lastindex)}
        return message_template.format(**extracted)
    except Exception as e:
        print(f"⚠️ Error extracting text: {str(e)}")
        return None

@timed_function
def get_address(yaml_data, extracted_text, theSection):
    extracted_text = extracted_text or ""
    address = "Address not found"
    address_heading = "Address Heading not found"
    message_if_identifier_found = "None"
    section_content = "None"
    for doc_section in yaml_data['docs']:
        # Process all questions including address and sections
        for question in doc_section.get('questions', []):
            # Check if identifier exists in text
            identifier = doc_section.get('identifier', '')
            if identifier and identifier in extracted_text:          
                message_if_identifier_found = doc_section['message_if_identifier_found']

            # Handle address specifically
            if 'address' in question:
                print(f"🔍 Processing address with pattern: {question['search_pattern']}")
                # Add centered address heading
                address_heading = question['address']                    
                if question.get('search_pattern') and question.get('extract_text', False):
                    address = extract_matching_text(
                        extracted_text,
                        question['search_pattern'],
                        question['extract_pattern'],
                        question['message_template']
                    )
            # Process all other sections
            if 'sections' in question:
                for section in question['sections']:
                    if section['section'] == theSection:
                        section_content = extract_matching_text(
                            extracted_text,
                            section['search_pattern'],
                            section['extract_pattern'],
                            section['message_template']
                        )

    return message_if_identifier_found, address_heading, address, section_content

def is_date_one_year_older(date_str):
    try:
        input_date = datetime.strptime(date_str, "%d-%b-%Y").date()
        today = datetime.now().date()
        one_year_ago = today - relativedelta(years=1)
        return input_date <= one_year_ago
    except ValueError:
        return False

@timed_function
def get_section(yaml_data, extracted_text, theSection):
    extracted_text = extracted_text or ""
    for doc_section in yaml_data['docs']:
        # Process all questions including address and sections
        for question in doc_section.get('questions', []):
            # Process all other sections
            if 'sections' in question:
                for section in question['sections']:
                    if section['section'] == theSection:
                        content = extract_matching_text(
                            extracted_text,
                            section['search_pattern'],
                            section['extract_pattern'],
                            section['message_template']
                        )
                        return content, section['message_if_none']
    return "None", f"Section {theSection} not found"

@timed_function
def process_zip(zip_path, output_docx, yaml_path):
    """Process ZIP file with improved error handling."""
    try:
        output_folder = "output_files/unzipped_files"
        os.makedirs(output_folder, exist_ok=True)
        
        yaml_data = load_yaml(yaml_path)
        doc = Document()
        
        # Add title and scope ONLY ONCE at the beginning
        doc.add_heading(yaml_data['general']['title'], level=0)
        scope = yaml_data['general']['scope'][0]
        doc.add_heading(scope['heading'], level=1)
        doc.add_paragraph(scope['body'])
        
        with zipfile.ZipFile(zip_path, 'r') as zip_ref:
            zip_ref.extractall(output_folder)
        
        # First collect all extracted text and check for identifiers
        all_extracted_text = []
        doc_identifiers = {doc_section['identifier']: doc_section 
                          for doc_section in yaml_data['docs'] 
                          if 'identifier' in doc_section}
        
        # Track which identifiers we've found
        found_identifiers = set()
        
        for file_name in os.listdir(output_folder):
            file_path = os.path.join(output_folder, file_name)
            
            try:
                if file_name.endswith(".pdf"):
                    extracted_text = extract_text_from_pdf(file_path)
                elif file_name.endswith(".docx"):
                    extracted_text = extract_text_from_docx(file_path)
                else:
                    continue
                
                if not extracted_text.strip():
                    continue
                
                # Check if this file contains any of our identifiers
                for identifier, doc_section in doc_identifiers.items():
                    if identifier in extracted_text:
                        found_identifiers.add(identifier)
                        all_extracted_text.append(extracted_text)
                        print(f"✅ Found identifier '{identifier}' in {file_name}")
                        break
                
            except Exception as e:
                print(f"⚠️ Error processing {file_name}: {str(e)}")
                continue
        
        if not all_extracted_text:
            print("❌ No files with matching identifiers found")
            doc.add_paragraph("No matching documents found with the required identifiers.")
            doc.save(output_docx)
            return
        
        # Combine all text for processing
        combined_text = "\n".join(all_extracted_text)
        
        # Save combined text for potential later use
        write_combined_text(combined_text)
        
        os.makedirs(os.path.dirname(output_docx), exist_ok=True)
        doc.save(output_docx)
        print(f"✅ Report generated: {output_docx}")
        
    except Exception as e:
        print(f"❌ Critical error processing ZIP: {str(e)}")
        raise

@timed_function
def process_section_groups(yaml_data, combined_text, doc):
    section_groups = [
        {
            "sections": ["Planning Permission", "Listed Building", "Conservation Area"],
            "all_none_message": "There are no planning permissions and no listed building or conservation consents."
        },
        {
            "sections": ["Contaminated Land", "Radon Gas"],
            "all_none_message": "The search shows a clear result with regards to contaminated land and radon gas in the area."
        },
        {
            "sections": ["Enforcement Notice", "Stop Notice"],
            "all_none_message": "There are no notices, orders, directions and proceedings under planning acts registered."
        },
        {
            "sections": ["Permanent stopping", "Waiting or loading restrictions"],
            "all_none_message": "There are no traffic, road or railway schemes registered."
        },
        {
            "sections": ["Drainage Agreement", "Drainage Consents"],
            "all_none_message": ("There are no drainage agreements or consents existing in relation to the property. "
                                "It would be prudent for you to acquire a drainage and water search to verify how the "
                                "drainage system of the property is managed. If the drains and sewers are maintained "
                                "privately, you may be required to maintain them.")
        }
    ]

    for group in section_groups:
        all_none = process_section_group(
            group["sections"], 
            yaml_data, 
            combined_text, 
            doc
        )
        if all_none:
            doc.add_paragraph(group["all_none_message"], style="List Bullet")
        else:
            process_sections(yaml_data, combined_text, doc, sections_to_process)

@timed_function
def process_section_group(sections, yaml_data, text, doc):
    all_none = True
    for section in sections:
        content, message_if_none = get_section(yaml_data, text, section)
        
        if content is not None:
            content = str(content).strip().rstrip(';:,.')
            
        if content and content.upper() not in ["NO", "NONE", "NOT APPLICABLE", ""]:
            all_none = False

    return all_none

@timed_function
def process_sections(yaml_data, combined_text, doc, sections_to_process):
    for section in sections_to_process:
        content, message_if_none = get_section(yaml_data, combined_text, section)
        
        if content is None or str(content).strip().upper() in ["NO", "NONE", "NOT APPLICABLE"]:
            doc.add_paragraph(message_if_none, style="List Bullet")
        else:
            # Special handling for Certificate of Lawfulness
            if section == "Certificate of Lawfulness" and "No Decision to date" in content:
                message = ". However, at the date of the search a decision had not yet been made. " \
                          "It is imperative that you contact the local council to ensure that " \
                          "the existing use of the property is lawful as you may be held " \
                          "liable if the property's use or development is unlawful"
                content = content + message
            doc.add_paragraph(content, style="List Bullet")

def test_section_config(section_config):
    test_cases = [
        ("1.1(a) A Planning Permission;\nNone", "There are no planning permissions."),
        ("Planning permission\n1.01(a)Where applicable", "There are no planning permissions."),
        ("1.1(a) Granted planning permission", "Granted")
    ]
    
    for text, expected in test_cases:
        assert process_section(text, section_config) == expected

def process_section(text, section_config):
    """
    Generic section processor that uses only YAML configuration
    """
    # First try the configured search pattern
    if match := re.search(
        section_config['search_pattern'], 
        text, 
        re.IGNORECASE | re.MULTILINE | re.DOTALL
    ):
        extracted = {f"extracted_text_{i+1}": match.group(i+1) 
                   for i in range(match.lastindex)}
        test_section_config(section_config)
        if (section_config.get('ambiguous_handling') and 
            not any(re.search(p, text) for p in section_config['detection_rules']['positive_indicators']) and
            not any(re.search(n, text) for n in section_config['detection_rules']['negative_indicators'])):
            
            if section_config['ambiguous_handling']['require_human_review']:
                flag_for_review(text)
            return section_config['ambiguous_handling']['default_message']

        # Check for negative indicators if they exist
        if 'detection_rules' in section_config:
            if any(
                re.search(indicator, match.group(0), re.IGNORECASE)
                for indicator in section_config['detection_rules']['negative_indicators']
            ):
                return section_config['message_if_none']
        
        return section_config['message_template'].format(**extracted)
    
    return section_config['message_if_none']

# Usage example:
yaml_data = load_yaml("config.yaml")
planning_config = next(
    s for s in yaml_data['docs'][0]['sections'] 
    if s['section'] == "Planning Permission"
)

text = "1.1(a) A Planning Permission;\nNone"
result = process_section(text, planning_config)
print(result)  # "There are no planning permissions."

# Main execution
if __name__ == "__main__":
    input_folder = "input_files"
    yaml_config = "config.yaml"
    output_file = "output_files/processed_doc.docx"
    
    
    # Otherwise proceed with normal ZIP processing
    zip_file_path = next(
        (os.path.join(input_folder, f) for f in os.listdir(input_folder) if f.endswith(".zip")),
        None
    )
    
    if zip_file_path:
        print(f"📂 Found ZIP file: {zip_file_path}")
        process_zip(zip_file_path, output_file, yaml_config)
    else:
        print("❌ No ZIP file found in 'input_files' folder.")

    # Check if we should process from existing combined text
    if os.path.exists(os.path.join("work_files", "combined_text.txt")):
        print("📄 Found existing combined_text.txt")
        combined_text = read_combined_text()
        if combined_text:
            yaml_data = load_yaml(yaml_config)
            doc = Document()

            # Headings
            heading = doc.add_heading(yaml_data['general']['title'], level=0)
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            message_if_identifier_found, address_heading, address, date_of_search = get_address(yaml_data, combined_text, "Search Date")
            address_heading = doc.add_heading(address_heading, level=2)
            address_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            para = add_formatted_paragraph(doc, address, italic=True)
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            scope = yaml_data['general']['scope'][0]
            doc.add_heading(scope['heading'], level=1)
            doc.add_paragraph(scope['body'])
            doc.add_paragraph("Local Authority Search", style="Heading 2")
            para = doc.add_paragraph(message_if_identifier_found)
            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            if is_date_one_year_older(date_of_search):
                message = f"The Search result date is {date_of_search}. This is not an up-to-date search result therefore any policies or permissions that may have been registered after that date will not be reflected on the said local authority search. We would advise you to acquire a new local search to acquire information that is up to date."
                doc.add_paragraph(message, style="List Bullet")

            # Process sections
            sections_to_process = [
                "Building Regulations",
                "Certificate of Lawfulness",
                "Highways",
                "Adoption Agreement",
                "Land required for Public Purposes",
                "Infringement of Building Regulations"
            ]
            process_sections(yaml_data, combined_text, doc, sections_to_process)

            # Process groups to check if all sections are None
            process_section_groups(yaml_data, combined_text, doc)
            doc.save(output_file)
            print(f"✅ Report generated from combined text: {output_file}")
            exit()
